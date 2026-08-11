"""
从扩充后的题目_expanded.docx 生成 data.js 和 problem_cache.js
v4: 每章15题（1模板+3绿+5青+3蓝+3紫），全部唯一无跨章重复
"""
import os, json, re
from docx import Document
from collections import Counter

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(BASE_DIR, '题目_expanded.docx')
DATA_PATH = os.path.join(BASE_DIR, 'js', 'data.js')
CACHE_PATH = os.path.join(BASE_DIR, 'js', 'problem_cache.js')

DIFF_MAP = {'模板': 1, '绿': 4, '青': 5, '蓝': 6, '紫': 7}

# 读取 data.js 获取章节结构
with open(DATA_PATH, 'r', encoding='utf-8') as f:
    old_data = f.read()
json_str = old_data.replace('const CHAPTERS =', '').strip().rstrip(';').strip()
json_str = json_str.replace("'", '"')
json_str = re.sub(r'([\{\,\[]\s*)([a-zA-Z_][a-zA-Z0-9_]*)\s*:', r'\1"\2":', json_str)
old_chapters = json.loads(json_str)

num_to_ch = {}
for ch in old_chapters:
    m = re.match(r'(\d+\.\d+(?:\.\d+)?)', ch['title'])
    if m:
        num_to_ch[m.group(1)] = ch

# 读取扩充后的 docx
doc = Document(DOCX_PATH)
titles = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
assert len(doc.tables) >= len(titles)

table_to_ch = {}
for i, title in enumerate(titles):
    m = re.match(r'Part (\d+\.\d+(?:\.\d+)?)', title)
    if m and m.group(1) in num_to_ch:
        table_to_ch[i] = num_to_ch[m.group(1)]

print(f"已匹配 {len(table_to_ch)} 个章节的表格")

# 解析每个章节的15道题
chapter_problems = {}  # ch_id -> [(pid, title, diff), ...]
all_pids = []

for ti in range(len(doc.tables)):
    ch = table_to_ch.get(ti)
    if ch is None:
        continue
    ch_id = ch['id']
    problems = []
    for row in doc.tables[ti].rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue
        diff_label = cells[0]
        pid = cells[1]
        ptitle = cells[2]
        if not pid.startswith('P'):
            continue
        if pid == 'P0000':
            print(f"  ⚠ {ch['title']}: 待补充题号")
            continue
        diff = DIFF_MAP.get(diff_label, 3)
        problems.append((pid, ptitle, diff, diff_label))
    chapter_problems[ch_id] = problems
    all_pids.extend(p[0] for p in problems)

# 统计
unique_pids = set(all_pids)
dup_pids = [p for p, c in Counter(all_pids).items() if c > 1]
print(f"\n总题目槽位: {len(all_pids)}")
print(f"唯一PID: {len(unique_pids)}")
if dup_pids:
    print(f"重复PID: {len(dup_pids)} 个")
else:
    print("全部唯一，无跨章重复 ✓")

# 章节统计
print(f"\n=== 章节统计 ===")
for ch in old_chapters:
    probs = chapter_problems.get(ch['id'], [])
    n = len(probs)
    flag = "✓" if n == 15 else f"⚠ ({n})"
    if n != 15:
        print(f"  {flag} {ch['title']}: {n} 题")
    else:
        print(f"  {flag} {ch['title']}")

# 难度统计
diff_counts = Counter()
for ch in old_chapters:
    for pid, pt, d, dl in chapter_problems.get(ch['id'], []):
        diff_counts[d] += 1
labels = {1: '模板', 2: '红', 3: '橙', 4: '绿', 5: '青', 6: '蓝', 7: '紫'}
print(f"\n=== 难度分布 ===")
for d in sorted(diff_counts):
    print(f"  {labels.get(d, str(d))}({d}): {diff_counts[d]}")

# ===== 生成 data.js =====
new_chapters = []
cache_entries = []

for ch in old_chapters:
    ch_id = ch['id']
    problems = chapter_problems.get(ch_id, [])
    
    modules = [ch['modules'][0]]  # 保留第一个 intro 模块
    
    for idx, (pid, ptitle, diff, dl) in enumerate(problems[:15], 1):
        modules.append({
            'id': f"{ch_id}_prob_{idx}",
            'title': f"第{idx}题",
            'type': 'problem',
            'luoguId': pid
        })
        cache_entries.append({'pid': pid, 'title': ptitle, 'diff': diff})
    
    new_chapters.append({
        'id': ch['id'], 'title': ch['title'],
        'icon': ch.get('icon', ''), 'description': ch.get('description', ''),
        'modules': modules
    })

data_js = 'const CHAPTERS = ' + json.dumps(new_chapters, ensure_ascii=False, indent=2) + ';\n'
with open(DATA_PATH, 'w', encoding='utf-8') as f:
    f.write(data_js)
print(f"\n已生成 data.js: {len(new_chapters)} 章节")

# ===== 生成 problem_cache.js (基础模板) =====
lines = ['// 洛谷题目本地缓存（从题目_expanded.docx 生成）\n', 'window.PROBLEM_CACHE = {\n']
seen_cache = set()
for entry in cache_entries:
    pid = entry['pid']
    if pid in seen_cache:
        continue
    seen_cache.add(pid)
    title = entry['title'].replace('\\', '\\\\').replace("'", "\\'")
    diff = entry['diff']
    lines.append(f"  '{pid}': {{\n")
    lines.append(f"    id: '{pid}',\n")
    lines.append(f"    title: '{title}',\n")
    lines.append(f"    difficulty: {diff},\n")
    lines.append(f"    description: '{title}\\\\n\\\\n**提示：** 题目将从洛谷加载完整描述。',\n")
    lines.append(f"    samples: [],\n")
    lines.append(f"    constraints: '详见洛谷题目页面',\n")
    lines.append(f"  }},\n")
lines[-1] = lines[-1].rstrip(',\n') + '\n'
lines.append('};\n')
with open(CACHE_PATH, 'w', encoding='utf-8') as f:
    f.writelines(lines)

print(f"已生成 problem_cache.js: {len(seen_cache)} 道题")
print("\n完成！接下来运行 batch_fetch.py 爬取完整题目数据")
print("然后运行 update_cache.py 更新缓存")
