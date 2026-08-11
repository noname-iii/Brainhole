"""
v3: 两阶段分配
  Pass1: 唯一分配（每个题号只给第一个声明它的章节）
  Pass2: 不足10题的章节从文档自有表补充（允许跨章重复，优先青/蓝/紫≥5）
"""
import os
import json, re, os
from docx import Document
from collections import Counter

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(BASE_DIR, '题目.docx')
DATA_PATH = os.path.join(BASE_DIR, 'js', 'data.js')
CACHE_PATH = os.path.join(BASE_DIR, 'js', 'problem_cache.js')

DIFF_MAP = {'模板': 1, '绿': 4, '青': 5, '蓝': 6, '紫': 7}

with open(DATA_PATH, 'r', encoding='utf-8') as f:
    old_data = f.read()
json_str = old_data.replace('const CHAPTERS =', '').strip().rstrip(';').strip()
json_str = json_str.replace("'", '"')
json_str = re.sub(r'([\{\,\[]\s*)([a-zA-Z_][a-zA-Z0-9_]*)\s*:', r'\1"\2":', json_str)
old_chapters = json.loads(json_str)

num_to_ch = {}
for ch in old_chapters:
    m = re.match(r'(\d+\.\d+(?:\.\d+)?)', ch['title'])
    if m:
        num_to_ch[m.group(1)] = ch

doc = Document(DOCX_PATH)
titles = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
assert len(doc.tables) == len(titles)

table_to_ch = {}
for i, title in enumerate(titles):
    m = re.match(r'Part (\d+\.\d+(?:\.\d+)?)', title)
    if m and m.group(1) in num_to_ch:
        table_to_ch[i] = num_to_ch[m.group(1)]

# === 解析每个章节在文档中的全部题目（含重复） ===
all_table_problems = {}  # ch_id → [(pid, title, diff, diff_label), ...] ordered as in table
for ti in range(len(doc.tables)):
    ch = table_to_ch.get(ti)
    if ch is None:
        continue
    ch_id = ch['id']
    problems = []
    for row in doc.tables[ti].rows[1:]:
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 3:
            continue
        diff_label = cells[0]
        pid = cells[1]
        ptitle = cells[2]
        if not pid.startswith('P'):
            continue
        diff = DIFF_MAP.get(diff_label, 3)
        problems.append((pid, ptitle, diff, diff_label))
    all_table_problems[ch_id] = problems

# === Pass 1: 唯一分配 ===
global_unique = set()
chapter_unique = {}  # ch_id → [(pid, title, diff), ...]
for ch in old_chapters:
    chapter_unique[ch['id']] = []

for ti in range(len(doc.tables)):
    ch = table_to_ch.get(ti)
    if ch is None:
        continue
    ch_id = ch['id']
    for pid, ptitle, diff, dl in all_table_problems[ch_id]:
        if pid not in global_unique:
            global_unique.add(pid)
            chapter_unique[ch_id].append((pid, ptitle, diff))

print(f"Pass 1: {len(global_unique)} 唯一题分配完成")

# === Pass 2: 补充不足10题的章节 ===
short_count = 0
for ch in old_chapters:
    ch_id = ch['id']
    have = len(chapter_unique[ch_id])
    if have >= 10:
        continue
    
    need = 10 - have
    
    # 从该章节文档表中取未被本小节已用的题目
    used_in_ch = set(p[0] for p in chapter_unique[ch_id])
    remaining = [(pid, ptitle, diff, dl) for pid, ptitle, diff, dl in all_table_problems[ch_id]
                 if pid not in used_in_ch]
    
    # 排序：优先青(5)+蓝(6)+紫(7)，然后绿(4)，最后模板(1)/其他
    def sort_key(x):
        d = x[2]
        if d >= 5:
            return (0, -d)  # diff越高越优先
        elif d == 4:
            return (1, 0)
        else:
            return (2, -d)
    
    remaining.sort(key=sort_key)
    
    supplement = remaining[:need]
    for pid, ptitle, diff, dl in supplement:
        chapter_unique[ch_id].append((pid, ptitle, diff))
    
    short_count += 1
    print(f"  +{len(supplement)} to {ch['title']} (now {len(chapter_unique[ch_id])})")

# === 统计 ===
print(f"\n=== 最终章节题目数 ===")
still_short = 0
for ch in old_chapters:
    n = len(chapter_unique[ch['id']])
    flag = "✓" if n >= 10 else "⚠"
    if n < 10:
        still_short += 1
        print(f"  {flag} {ch['title']}: {n} 题")
    else:
        print(f"  {flag} {ch['title']}: {n} 题")

total_pids = sum(len(v) for v in chapter_unique.values())
print(f"\n总计: {total_pids} 题(槽位)")

# === 生成 data.js ===
new_chapters = []
cache_entries = []
for ch in old_chapters:
    ch_id = ch['id']
    problems = chapter_unique[ch_id]
    
    modules = [ch['modules'][0]]
    taken = problems[:10]
    
    for idx, (pid, ptitle, diff) in enumerate(taken, 1):
        modules.append({
            'id': f"{ch_id}_prob_{idx}",
            'title': f"第{idx}题",
            'type': 'problem',
            'luoguId': pid
        })
        cache_entries.append({'pid': pid, 'title': ptitle, 'diff': diff})
    
    new_chapters.append({
        'id': ch['id'], 'title': ch['title'],
        'icon': ch.get('icon', ''), 'description': ch.get('description', ''),
        'modules': modules
    })

data_js = 'const CHAPTERS = ' + json.dumps(new_chapters, ensure_ascii=False, indent=2) + ';\n'
with open(DATA_PATH, 'w', encoding='utf-8') as f:
    f.write(data_js)

# === 生成 problem_cache.js ===
lines = ['// 洛谷题目本地缓存（从题目.docx 生成）\n', 'window.PROBLEM_CACHE = {\n']
seen_cache = set()
for entry in cache_entries:
    pid = entry['pid']
    if pid in seen_cache:
        continue
    seen_cache.add(pid)
    title = entry['title'].replace('\\', '\\\\').replace("'", "\\'")
    diff = entry['diff']
    lines.append(f"  '{pid}': {{\n")
    lines.append(f"    id: '{pid}',\n")
    lines.append(f"    title: '{title}',\n")
    lines.append(f"    difficulty: {diff},\n")
    lines.append(f"    description: '{title}\\\\n\\\\n**提示：** 若网络可用，将自动从洛谷加载完整题目描述。',\n")
    lines.append(f"    samples: [],\n")
    lines.append(f"    constraints: '详见洛谷题目页面',\n")
    lines.append(f"  }},\n")
lines[-1] = lines[-1].rstrip(',\n') + '\n'
lines.append('};\n')
with open(CACHE_PATH, 'w', encoding='utf-8') as f:
    f.writelines(lines)

# === 最终统计 ===
all_pids = [e['pid'] for e in cache_entries]
uniq_pids = set(all_pids)
dup_pids = [p for p, c in Counter(all_pids).items() if c > 1]

print(f"\n=== 最终统计 ===")
print(f"data.js: {len(cache_entries)} 题槽位")
print(f"唯一题目: {len(uniq_pids)}")
print(f"跨章重复({len(dup_pids)}个): {dup_pids[:20]}")

diff_counts = Counter(e['diff'] for e in cache_entries)
labels = {0:'?', 1:'红', 2:'橙', 3:'黄', 4:'绿', 5:'青', 6:'蓝', 7:'紫'}
for d in sorted(diff_counts):
    print(f"  {labels.get(d, str(d))}: {diff_counts[d]}")
